"""Exportação do prontuário para Word (.docx) e PDF."""

from __future__ import annotations

from pathlib import Path

TITLE = "Prontuário de Consulta"


def export_docx(path: Path, body: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(TITLE, level=1)
    for paragraph in body.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(path))


def export_pdf(path: Path, body: str) -> None:
    from xml.sax.saxutils import escape

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=TITLE,
    )
    story = [Paragraph(TITLE, styles["Title"]), Spacer(1, 4 * mm)]
    for paragraph in body.split("\n"):
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph), styles["BodyText"]))
        else:
            story.append(Spacer(1, 3 * mm))
    doc.build(story)
